"""文件处理服务"""
import io
import warnings
import zipfile
from pathlib import Path
from typing import Any, List, Tuple

import fitz  # pymupdf
from docx import Document
from openpyxl import load_workbook
from PIL import Image
from backend.services.extraction_utils import read_excel_as_rows, rows_to_text

with warnings.catch_warnings():
    warnings.filterwarnings(
        "ignore",
        message=r"PyPDF2 is deprecated\. Please move to the pypdf library instead\.",
        category=DeprecationWarning,
    )
    from PyPDF2 import PdfReader


class FileService:
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """获取文件类型"""
        suffix = Path(filename).suffix.lower()
        type_map = {
            ".pdf": "pdf",
            ".png": "image",
            ".jpg": "image",
            ".jpeg": "image",
            ".xlsx": "excel",
            ".xls": "excel",
            ".docx": "word",
            ".doc": "word",
        }
        return type_map.get(suffix, "unknown")
    
    @staticmethod
    def read_pdf_text(file_bytes: bytes) -> str:
        """读取 PDF 文本（仅限文本型 PDF）
        
        尝试从 PDF 中直接提取文本内容。
        
        Args:
            file_bytes: PDF 文件的字节数据
            
        Returns:
            提取的文本内容。如果 PDF 是扫描件（无可提取文本），返回空字符串。
            如果读取失败，返回以 "[PDF_ERROR]" 开头的错误信息。
            
        **Validates: Requirement 3.2**
        
        WHEN a PDF file is uploaded, THE File_Service SHALL first attempt 
        to extract text directly
        """
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    # 清理提取的文本（去除多余空白）
                    cleaned_text = text.strip()
                    if cleaned_text:
                        text_parts.append(cleaned_text)
            
            result = "\n".join(text_parts)
            return result
            
        except Exception as e:
            # 使用特定前缀标识错误，便于调用方识别
            return f"[PDF_ERROR] 读取失败: {str(e)}"
    
    @staticmethod
    def is_pdf_text_valid(content: str) -> bool:
        """检查 PDF 提取的文本是否有效
        
        判断 PDF 文本提取结果是否有效（非空、非错误、有实际内容）。
        用于判断 PDF 是否为扫描件。
        
        Args:
            content: read_pdf_text() 返回的内容
            
        Returns:
            True 如果文本有效（可用于 AI 提取），False 如果无效（可能是扫描件）
            
        **Validates: Requirement 3.3**
        
        IF direct PDF text extraction yields empty or invalid content, 
        THEN THE Upload_Module SHALL prompt user to convert PDF to images
        """
        # 检查是否为 None 或空
        if content is None:
            return False
        
        # 检查是否为错误信息
        if content.startswith("[PDF_ERROR]"):
            return False
        
        # 检查是否为空或仅包含空白字符
        if not content or not content.strip():
            return False
        
        # 检查是否有足够的有效字符（排除仅有少量乱码的情况）
        # 至少需要 10 个非空白字符才认为是有效文本
        non_whitespace = ''.join(content.split())
        if len(non_whitespace) < 10:
            return False
        
        return True
    
    @staticmethod
    def read_excel(file_bytes: bytes, sheet_name: str = None, read_all_sheets: bool = True) -> str:
        """读取 Excel 文件所有单元格内容
        
        从 Excel 文件中读取单元格内容，转换为文本格式。
        默认读取所有工作表，也可指定特定工作表。
        
        **智能处理公式**：同时读取 data_only 和非 data_only 模式，
        优先使用计算后的值，如果为空则使用公式文本或原始值。
        
        Args:
            file_bytes: Excel 文件的字节数据
            sheet_name: 可选，指定要读取的工作表名称。
                       如果指定，则只读取该工作表。
            read_all_sheets: 是否读取所有工作表，默认 True。
                            当 sheet_name 指定时此参数被忽略。
            
        Returns:
            包含所有非空单元格内容的文本字符串。
            每行单元格用 " | " 分隔，行与行之间用换行符分隔。
            多个工作表之间用 "=== Sheet: xxx ===" 分隔。
            如果读取失败，返回以 "[Excel 读取失败:" 开头的错误信息。
            
        **Validates: Requirement 3.4**
        
        WHEN an Excel file is uploaded, THE File_Service SHALL read all 
        cell contents as text
        """
        try:
            # 同时加载两种模式，以便合并结果
            # data_only=True: 读取公式的计算结果（如果 Excel 曾被打开计算过）
            # data_only=False: 读取公式文本和原始值
            wb_data = load_workbook(io.BytesIO(file_bytes), data_only=True)
            wb_formula = load_workbook(io.BytesIO(file_bytes), data_only=False)
            
            # 确定要读取的工作表列表
            if sheet_name:
                if sheet_name not in wb_data.sheetnames:
                    return f"[Excel 读取失败: 工作表 '{sheet_name}' 不存在]"
                sheets_to_read = [sheet_name]
            elif read_all_sheets:
                sheets_to_read = wb_data.sheetnames
            else:
                sheets_to_read = [wb_data.active.title]
            
            all_content = []
            
            for sname in sheets_to_read:
                ws_data = wb_data[sname]
                ws_formula = wb_formula[sname]
                
                rows = []
                
                # 获取工作表的实际使用范围
                max_row = max(ws_data.max_row or 1, ws_formula.max_row or 1)
                max_col = max(ws_data.max_column or 1, ws_formula.max_column or 1)
                
                for row_idx in range(1, max_row + 1):
                    row_data = []
                    
                    for col_idx in range(1, max_col + 1):
                        # 获取 data_only 模式的值（计算结果）
                        cell_data = ws_data.cell(row=row_idx, column=col_idx).value
                        # 获取非 data_only 模式的值（公式或原始值）
                        cell_formula = ws_formula.cell(row=row_idx, column=col_idx).value
                        
                        # 智能选择：优先使用计算结果，否则使用公式/原始值
                        if cell_data is not None:
                            row_data.append(str(cell_data))
                        elif cell_formula is not None:
                            # 如果是公式（以 = 开头），标记一下
                            cell_str = str(cell_formula)
                            if cell_str.startswith('='):
                                # 公式未计算，保留公式文本供 AI 参考
                                row_data.append(f"[公式]{cell_str}")
                            else:
                                row_data.append(cell_str)
                        else:
                            row_data.append("")
                    
                    # 只保留有内容的行
                    if any(cell.strip() for cell in row_data):
                        rows.append(" | ".join(row_data))
                
                if rows:
                    if len(sheets_to_read) > 1:
                        all_content.append(f"=== Sheet: {sname} ===")
                    all_content.extend(rows)
            
            return "\n".join(all_content)
            
        except Exception as e:
            return f"[Excel 读取失败: {str(e)}]"
    
    @staticmethod
    def get_excel_sheets(file_bytes: bytes) -> list:
        """获取 Excel 所有 sheet 名称"""
        try:
            wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
            return wb.sheetnames
        except Exception:
            return []
    
    @staticmethod
    def read_word(file_bytes: bytes) -> str:
        """读取 Word 文档"""
        try:
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            tables: list[str] = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                    if cells:
                        tables.append(" | ".join(cells))
            return "\n".join(paragraphs + tables)
        except Exception as e:
            return f"[Word 读取失败: {str(e)}]"

    @staticmethod
    def extract_word_images(file_bytes: bytes) -> list[bytes]:
        """Extract embedded image bytes from a DOCX package."""
        images: list[bytes] = []
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
                for member_name in archive.namelist():
                    lowered = member_name.lower()
                    if not lowered.startswith("word/media/"):
                        continue
                    if not lowered.endswith((".png", ".jpg", ".jpeg", ".bmp", ".webp")):
                        continue
                    try:
                        images.append(archive.read(member_name))
                    except Exception:
                        continue
        except Exception:
            return []
        return images

    @staticmethod
    def read_excel_rows(file_bytes: bytes) -> list[dict[str, str]]:
        """Read Excel rows without flattening table structure."""
        try:
            return read_excel_as_rows(file_bytes)
        except Exception:
            return []

    @staticmethod
    def extract_content(file_bytes: bytes, file_type: str, *, filename: str = "") -> dict[str, Any]:
        """Unified extraction entry returning text and optional row data."""
        if file_type == "pdf":
            return {"text": FileService.read_pdf_text(file_bytes), "rows": []}
        if file_type == "excel":
            rows = FileService.read_excel_rows(file_bytes)
            return {"text": rows_to_text(rows), "rows": rows}
        if file_type == "word":
            return {"text": FileService.read_word(file_bytes), "rows": []}
        if file_type == "image":
            return {"text": "", "rows": []}
        raise ValueError(f"Unsupported file type for content extraction: {file_type} ({filename})")

    @staticmethod
    def extract_text(file_bytes: bytes, file_type: str, *, filename: str = "") -> str:
        """统一的文本提取入口，供 PDF / DOCX / XLSX 等文件复用。"""
        return FileService.extract_content(file_bytes, file_type, filename=filename).get("text", "")
    
    @staticmethod
    def image_to_bytes(uploaded_file) -> bytes:
        """将上传的图片转为 bytes"""
        return uploaded_file.read()
    
    @staticmethod
    def compress_image(image_bytes: bytes, max_size: int = 4 * 1024 * 1024) -> bytes:
        """压缩图片（百度 OCR 限制 4MB）
        
        Args:
            image_bytes: 原始图片字节数据
            max_size: 最大允许大小，默认 4MB（百度 OCR 限制）
            
        Returns:
            压缩后的图片字节数据，如果原图小于限制则返回原图
            
        **Validates: Requirement 2.5**
        
        IF the file is an image larger than 4MB, THEN THE File_Service 
        SHALL compress it before OCR processing
        """
        # 如果图片已经小于限制，直接返回
        if len(image_bytes) <= max_size:
            return image_bytes
        
        try:
            img = Image.open(io.BytesIO(image_bytes))
            
            # 转换为 RGB 模式（处理 RGBA 或其他模式）
            if img.mode in ('RGBA', 'P', 'LA'):
                img = img.convert('RGB')
            
            # 逐步降低质量
            for quality in [85, 70, 50, 30]:
                buffer = io.BytesIO()
                img.save(buffer, format="JPEG", quality=quality)
                if buffer.tell() <= max_size:
                    return buffer.getvalue()
            
            # 如果还是太大，缩小尺寸
            width, height = img.size
            while True:
                width = width // 2
                height = height // 2
                
                if width < 100 or height < 100:
                    # 尺寸太小了，返回当前最佳结果
                    break
                    
                resized_img = img.resize((width, height), Image.Resampling.LANCZOS)
                buffer = io.BytesIO()
                resized_img.save(buffer, format="JPEG", quality=50)
                
                if buffer.tell() <= max_size:
                    return buffer.getvalue()
            
            # 返回最后一次压缩的结果（即使超过限制）
            buffer = io.BytesIO()
            img.resize((width * 2, height * 2), Image.Resampling.LANCZOS).save(
                buffer, format="JPEG", quality=30
            )
            return buffer.getvalue()
            
        except Exception as e:
            # 如果压缩失败，返回原始图片
            return image_bytes
    
    @staticmethod
    def validate_file_size(file_bytes: bytes, max_mb: int = 200) -> bool:
        """验证文件大小是否在限制内
        
        Args:
            file_bytes: 文件内容的字节数据
            max_mb: 最大允许的文件大小（MB），默认 200MB
            
        Returns:
            True 如果文件大小在限制内，False 如果超过限制
            
        **Validates: Requirement 2.2**
        """
        size_in_mb = len(file_bytes) / (1024 * 1024)
        return size_in_mb <= max_mb
    
    @staticmethod
    def validate_file_extension(filename: str, allowed_extensions: list) -> bool:
        """验证文件扩展名是否在允许列表中
        
        Args:
            filename: 文件名（包含扩展名）
            allowed_extensions: 允许的扩展名列表（不含点号，如 ['pdf', 'png', 'jpg']）
            
        Returns:
            True 如果扩展名在允许列表中，False 否则
            
        **Validates: Requirement 2.1**
        
        WHEN a user uploads a file, THE Upload_Module SHALL validate the file 
        extension matches the allowed formats for the selected Data_Type
        """
        # 处理空文件名
        if not filename:
            return False
        
        # 提取扩展名（大小写不敏感）
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        
        # 处理无扩展名的情况
        if not ext:
            return False
        
        # 将允许列表转为小写进行比较
        allowed_lower = [e.lower() for e in allowed_extensions]
        
        return ext in allowed_lower

    @staticmethod
    def pdf_to_images(file_bytes: bytes, dpi: int = 200) -> List[bytes]:
        """将 PDF 转换为图片列表（用于扫描件 OCR）
        
        使用 pymupdf 将 PDF 每一页转换为高清图片。
        
        Args:
            file_bytes: PDF 文件的字节数据
            dpi: 输出图片的分辨率，默认 200（清晰且文件大小适中）
                 - 72: 低质量，可能模糊
                 - 150: 中等质量
                 - 200: 推荐，清晰且 OCR 效果好
                 - 300: 高质量，文件较大
            
        Returns:
            图片字节数据列表，每个元素对应 PDF 的一页（PNG 格式）
            如果转换失败，返回空列表
            
        Note:
            - 转换后的图片可能超过百度 OCR 的 4MB 限制，需要配合 compress_image 使用
            - 多页 PDF 会返回多张图片，需要逐一 OCR 后合并文字
        """
        try:
            # 打开 PDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            images = []
            
            # 计算缩放矩阵（dpi / 72，因为 PDF 默认 72 dpi）
            zoom = dpi / 72
            matrix = fitz.Matrix(zoom, zoom)
            
            # 遍历每一页
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # 渲染为像素图
                pix = page.get_pixmap(matrix=matrix)
                
                # 转换为 PNG 字节
                img_bytes = pix.tobytes("png")
                images.append(img_bytes)
            
            doc.close()
            return images
            
        except Exception as e:
            # 转换失败，返回空列表
            return []
    
    @staticmethod
    def get_pdf_page_count(file_bytes: bytes) -> int:
        """获取 PDF 页数
        
        Args:
            file_bytes: PDF 文件的字节数据
            
        Returns:
            PDF 页数，如果读取失败返回 0
        """
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            count = len(doc)
            doc.close()
            return count
        except Exception:
            return 0
