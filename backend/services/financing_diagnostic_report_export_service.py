from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document


DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_UNAVAILABLE_MESSAGE = "PDF 导出依赖未配置，请先配置 PDF 渲染环境"


class FinancingDiagnosticReportExportError(RuntimeError):
    pass


class SnapshotNotFoundError(FinancingDiagnosticReportExportError):
    pass


class PdfExportUnavailableError(FinancingDiagnosticReportExportError):
    pass


def _as_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def can_export_financing_diagnostic_report_snapshot(role: str | None) -> bool:
    return str(role or "").lower() in {"admin", "operator", "viewer"}


def _safe_filename_part(value: Any, fallback: str) -> str:
    text = str(value or "").strip() or fallback
    text = re.sub(r'[\\/:*?"<>|\r\n\t]+', "_", text)
    text = re.sub(r"\s+", "_", text)
    return text[:80] or fallback


def _customer_name(snapshot: dict[str, Any]) -> str:
    report_json = _as_dict(snapshot.get("report_json"))
    customer_summary = _as_dict(report_json.get("customer_summary"))
    return str(customer_summary.get("customer_name") or snapshot.get("customer_id") or "未命名客户")


def _snapshot_markdown(snapshot: dict[str, Any]) -> str:
    markdown = str(snapshot.get("report_markdown") or "").strip()
    if markdown:
        return markdown
    report_json = _as_dict(snapshot.get("report_json"))
    return str(report_json.get("report_markdown") or "# 客户融资诊断报告").strip()


def _add_markdown_line(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    if stripped.startswith("### "):
        document.add_heading(stripped[4:].strip(), level=3)
        return
    if stripped.startswith("## "):
        document.add_heading(stripped[3:].strip(), level=2)
        return
    if stripped.startswith("# "):
        document.add_heading(stripped[2:].strip(), level=1)
        return
    if stripped.startswith("- "):
        document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        return
    document.add_paragraph(stripped)


def _markdown_to_docx_bytes(markdown: str, snapshot: dict[str, Any]) -> bytes:
    document = Document()
    document.add_heading("客户融资诊断报告", level=0)
    document.add_paragraph(f"报告版本：{snapshot.get('report_version') or '未记录'}")
    document.add_paragraph(f"生成时间：{snapshot.get('generated_at') or '未记录'}")
    document.add_paragraph(f"生成人：{snapshot.get('generated_by') or '未记录'}")
    document.add_paragraph("说明：本报告基于已保存的历史快照生成，后续资料变化不会影响本报告内容。")
    document.add_paragraph("")

    for raw_line in markdown.splitlines():
        _add_markdown_line(document, raw_line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_export_filename(snapshot: dict[str, Any], extension: str) -> str:
    customer_name = _safe_filename_part(_customer_name(snapshot), "未命名客户")
    version = _safe_filename_part(snapshot.get("report_version"), "未记录版本")
    return f"融资诊断报告_{customer_name}_{version}.{extension.lstrip('.')}"


class FinancingDiagnosticReportExportService:
    def __init__(self, storage_service: Any) -> None:
        self.storage_service = storage_service

    async def _load_snapshot(self, customer_id: str, report_id: str) -> dict[str, Any]:
        getter = getattr(self.storage_service, "get_financing_diagnostic_report_snapshot", None)
        if not callable(getter):
            raise SnapshotNotFoundError("未找到该融资诊断报告快照")
        snapshot = await getter(customer_id, report_id)
        if not snapshot:
            raise SnapshotNotFoundError("未找到该融资诊断报告快照")
        return snapshot

    async def export_docx(self, customer_id: str, report_id: str) -> dict[str, Any]:
        snapshot = await self._load_snapshot(customer_id, report_id)
        content = _markdown_to_docx_bytes(_snapshot_markdown(snapshot), snapshot)
        return {
            "content": content,
            "filename": build_export_filename(snapshot, "docx"),
            "media_type": DOCX_MEDIA_TYPE,
        }

    async def export_pdf(self, customer_id: str, report_id: str) -> dict[str, Any]:
        await self._load_snapshot(customer_id, report_id)
        raise PdfExportUnavailableError(PDF_UNAVAILABLE_MESSAGE)
